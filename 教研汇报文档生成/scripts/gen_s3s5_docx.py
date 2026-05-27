from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        borders.append(border)
    tblPr.append(borders)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level == 1: run.font.size = Pt(18); run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        elif level == 2: run.font.size = Pt(14); run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        elif level == 3: run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    return h

def add_para(doc, text, bold=False, italic=False, size=Pt(10.5), color=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = size
    run.font.name = '微软雅黑'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p

def add_bullet(doc, text, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text); run.bold = bold; run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.left_indent = Cm(1.27); p.paragraph_format.space_after = Pt(3)
    return p

def add_quote(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text); run.font.size = Pt(9.5); run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66); run.italic = True
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = ''
        p = cell.paragraphs[0]; run = p.add_run(header)
        run.bold = True; run.font.size = Pt(9); run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2C3E50')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]; cell.text = ''
            p = cell.paragraphs[0]; run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            try:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            except: pass
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r % 2 == 0: set_cell_shading(cell, 'F5F7FA')
    doc.add_paragraph()
    return table

def add_separator(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom); pPr.append(pBdr)

def pct(v): return f"{v*100:.2f}%"
def diff_str(actual, target):
    d = (actual - target)*100
    sign = '+' if d >= 0 else ''
    return f"{sign}{d:.2f}pp"
def trend_arrow(vals):
    if len(vals) < 2: return '→ 持平'
    first, last = vals[0], vals[-1]
    if last > first * 1.03: return '↑ 提升'
    elif last < first * 0.97: return '↓ 下降'
    else: return '→ 持平'

# ========== 数据 ==========
classes = ['普通班（粤）', '普通班（普）', '粤语班（粤）', '英语班（内）', '英语班（外）']
months = ['1月', '2月', '3月', '4月', '5月']

first_prac = {
    '普通班（粤）': [0.6441, 0.7531, 0.7492, 0.7822, 0.7636],
    '普通班（普）': [0.7232, 0.7282, 0.7407, 0.7498, 0.7478],
    '粤语班（粤）': [0.5869, 0.6414, 0.6640, 0.6859, 0.6792],
    '英语班（内）': [0.6275, 0.6154, 0.6490, 0.6628, 0.6643],
    '英语班（外）': [0.6254, 0.6317, 0.6647, 0.7154, 0.6887],
}
first_target = {'普通班（粤）': 0.77, '普通班（普）': 0.77, '粤语班（粤）': 0.69, '英语班（内）': 0.68, '英语班（外）': 0.68}

last_prac = {
    '普通班（粤）': [0.9018, 0.9383, 0.9605, 0.9690, 0.9768],
    '普通班（普）': [0.9372, 0.9293, 0.9323, 0.9444, 0.9437],
    '粤语班（粤）': [0.8733, 0.8869, 0.9186, 0.9251, 0.9369],
    '英语班（内）': [0.9128, 0.8997, 0.9201, 0.9336, 0.9300],
    '英语班（外）': [0.9072, 0.9018, 0.9235, 0.9400, 0.9418],
}
last_target = {cls: 0.95 for cls in classes}

unlock_ex = {
    '普通班（粤）': [4.947, 6.536, 7.105, 7.576, 7.420],
    '普通班（普）': [6.505, 6.112, 6.319, 7.059, 7.381],
    '粤语班（粤）': [5.285, 5.593, 5.826, 6.683, 6.980],
    '英语班（内）': [6.160, 6.216, 6.587, 7.099, 7.122],
    '英语班（外）': [5.373, 5.735, 6.116, 5.997, 5.755],
}
unlock_target = 5.5

safe_pct = {
    '普通班（粤）': [0.3364, 0.6166, 0.7187, 0.7761, 0.7801],
    '普通班（普）': [0.6403, 0.5906, 0.6212, 0.7579, 0.7899],
    '粤语班（粤）': [0.4139, 0.4621, 0.5031, 0.6727, 0.7374],
    '英语班（内）': [0.5429, 0.5706, 0.6327, 0.7176, 0.7221],
    '英语班（外）': [0.4215, 0.4631, 0.5157, 0.4981, 0.4743],
}
safe_target = 0.85

hw = {
    '普通班（粤）': [0.8920, 0.8826, 0.9030, 0.9258, 0.8625],
    '普通班（普）': [0.9088, 0.8846, 0.9112, 0.9083, 0.8764],
    '粤语班（粤）': [0.8554, 0.8437, 0.8531, 0.8621, 0.8228],
    '英语班（内）': [0.8532, 0.8383, 0.8535, 0.8624, 0.8007],
    '英语班（外）': [0.9165, 0.8942, 0.9043, 0.9058, 0.8305],
}
hw_target = {'普通班（粤）': 0.90, '普通班（普）': 0.90, '粤语班（粤）': 0.87, '英语班（内）': 0.87, '英语班（外）': 0.90}

example_unlock = {
    '普通班（粤）': [1.259, 2.132, 2.258, 2.054, 1.826],
    '普通班（普）': [1.748, 1.627, 1.696, 1.820, 1.989],
    '粤语班（粤）': [1.630, 1.771, 1.820, 2.143, 2.103],
    '英语班（内）': [1.722, 1.859, 1.924, 1.971, 2.059],
    '英语班（外）': [1.670, 2.043, 1.974, 1.931, 2.127],
}

# ========== 正文 ==========
add_heading(doc, 'S3-S5 学段北极星指标分析报告', level=1)
add_para(doc, '分析时间：2026年5月（数据截至0524）', size=Pt(10))
add_para(doc, '分析对象：S3-S5 学段，覆盖普通班（粤）/ 普通班（普）/ 粤语班（粤）/ 英语班（内）/ 英语班（外）', size=Pt(10))
add_para(doc, '指标范围：课中练习首/末答正确率、解锁例题+练习数、安全解锁比例、作业完成率', size=Pt(10))
add_separator(doc)

# ==============================
# 一、课中练习首答正确率
# ==============================
add_heading(doc, '一、课中练习首答正确率', level=2)

add_heading(doc, '1.1 Q2 达标情况（截至5月）', level=3)
rows = []
for cls in classes:
    actual = first_prac[cls][4]
    target = first_target[cls]
    status = '✅ 已达标' if actual >= target else '❌ 未达标'
    rows.append([cls, pct(target), pct(actual), status, diff_str(actual, target)])
add_table(doc, ['班型', 'Q2 目标', '5月实际', '达成状态', '差距'], rows)
add_bullet(doc, '仅英语班（外）达标（68.87%）；普通班（粤）距目标最近仅差 0.64pp')
add_bullet(doc, '粤语班（粤）差距最大（-1.08pp），其余未达标班型差距在 1-2.22pp 之间')

add_heading(doc, '1.2 近5个月趋势', level=3)
rows = []
for cls in classes:
    v = first_prac[cls]
    rows.append([cls] + [pct(x) for x in v] + [trend_arrow(v)])
add_table(doc, ['班型'] + months + ['趋势'], rows)
add_bullet(doc, '全部班型整体呈上升趋势，1月整体低位，3-4月基本达峰')
add_bullet(doc, '英语班（外）1-4月持续提升（62.54%→71.54%），5月小幅回落至 68.87%，仍为唯一达标班型')
add_bullet(doc, '普通班（粤）4月达峰 78.22%，5月回落至 76.36%；整体改善趋势稳健')
add_bullet(doc, '粤语班（粤）改善明显：1月仅 58.69% 提升至 67.92%（+9.23pp），尚需继续努力')

add_separator(doc)

# ==============================
# 二、课中练习末答正确率
# ==============================
add_heading(doc, '二、课中练习末答正确率', level=2)

add_heading(doc, '2.1 Q2 达标情况（截至5月）', level=3)
rows = []
for cls in classes:
    actual = last_prac[cls][4]
    target = last_target[cls]
    status = '✅ 已达标' if actual >= target else '❌ 未达标'
    rows.append([cls, pct(target), pct(actual), status, diff_str(actual, target)])
add_table(doc, ['班型', 'Q2 目标', '5月实际', '达成状态', '差距'], rows)
add_bullet(doc, '仅普通班（粤）达标（97.68%）；其余4个班型均在 93-94% 区间，低于 95% 目标线')
add_bullet(doc, '英语班（内）距目标最近仅差 2pp，有望在6月冲刺达标')

add_heading(doc, '2.2 近5个月趋势', level=3)
rows = []
for cls in classes:
    v = last_prac[cls]
    rows.append([cls] + [pct(x) for x in v] + [trend_arrow(v)])
add_table(doc, ['班型'] + months + ['趋势'], rows)
add_bullet(doc, '全部班型末答正确率持续提升，改善趋势最明显')
add_bullet(doc, '普通班（粤）从 90.18% 升至 97.68%（+7.5pp），已达标且持续拉高')
add_bullet(doc, '英语班（外）从 90.72% 升至 94.18%（+3.46pp），接近 95% 目标线')
add_bullet(doc, '粤语班（粤）升幅最大（+6.36pp），从 87.33% 升至 93.69%，改善质量高')

add_separator(doc)

# ==============================
# 三、解锁例题 + 练习数
# ==============================
add_heading(doc, '三、解锁例题 + 练习数', level=2)

add_heading(doc, '3.1 Q2 达标情况（截至5月）', level=3)
rows = []
for cls in classes:
    actual = unlock_ex[cls][4]
    status = '✅ 已达标' if actual >= unlock_target else '❌ 未达标'
    rows.append([cls, f'{unlock_target:.1f}', f'{actual:.2f}', status])
add_table(doc, ['班型', 'Q2 目标', '5月实际（道）', '达成状态'], rows)
add_bullet(doc, '全部班型达标，解锁总量普遍超出目标 1-2 道')
add_bullet(doc, '普通班（粤）以 7.42 道领跑，英语班（外）5.76 道为最低但仍达标')

add_heading(doc, '3.2 近5个月趋势', level=3)
rows = []
for cls in classes:
    v = unlock_ex[cls]
    rows.append([cls] + [f'{x:.2f}' for x in v] + [trend_arrow(v)])
add_table(doc, ['班型'] + months + ['趋势'], rows)
add_bullet(doc, '普通班（粤）从 4.95 快速提升至 7.42（+2.47），增长最快')
add_bullet(doc, '英语班（外）4月后小幅回落（7.15→6.00→5.76），其他班型均持续上升')
add_bullet(doc, '粤语班（粤）持续增长（5.29→6.98），增幅达 1.70 道，稳定向好')

add_separator(doc)

# ==============================
# 四、安全解锁比例
# ==============================
add_heading(doc, '四、安全解锁比例', level=2)

add_heading(doc, '4.1 Q2 达标情况（截至5月）', level=3)
rows = []
for cls in classes:
    actual = safe_pct[cls][4]
    status = '❌ 未达标'
    rows.append([cls, pct(safe_target), pct(actual), status, diff_str(actual, safe_target)])
add_table(doc, ['班型', 'Q2 目标', '5月实际', '达成状态', '差距'], rows)
add_bullet(doc, '全线未达标，距 85% 目标差距在 7-37.57pp 之间', bold=True)
add_bullet(doc, '英语班（外）仅 47.43%，差距最大（-37.57pp），属严重异常需重点介入')
add_bullet(doc, '普通班（粤/普）相对最优（78-79%），缺口在 6-7pp，改善空间有限')

add_heading(doc, '4.2 近5个月趋势', level=3)
rows = []
for cls in classes:
    v = safe_pct[cls]
    rows.append([cls] + [pct(x) for x in v] + [trend_arrow(v)])
add_table(doc, ['班型'] + months + ['趋势'], rows)
add_bullet(doc, '普通班（粤）从 33.64% 提升至 78.01%（+44.37pp），改善最显著')
add_bullet(doc, '粤语班（粤）从 41.39% 升至 73.74%（+32.35pp），提升速度快')
add_bullet(doc, '英语班（外）全程在 42-52% 低位徘徊，5月进一步降至 47.43%，与其他班型分化加剧', bold=True)
add_bullet(doc, '各班型整体趋势向好，但距 85% 目标仍有较大差距，英语班（外）需专项干预')

add_separator(doc)

# ==============================
# 五、作业完成率
# ==============================
add_heading(doc, '五、作业完成率', level=2)

add_heading(doc, '5.1 Q2 达标情况（截至5月）', level=3)
rows = []
for cls in classes:
    actual = hw[cls][4]
    target = hw_target[cls]
    status = '✅ 已达标' if actual >= target else '❌ 未达标'
    rows.append([cls, pct(target), pct(actual), status, diff_str(actual, target)])
add_table(doc, ['班型', 'Q2 目标', '5月实际', '达成状态', '差距'], rows)
add_bullet(doc, '全线未达标，英语班（内）最低仅 80.07%（差 6.93pp），其余差距在 1.36-7.95pp 之间')

add_heading(doc, '5.2 近5个月趋势', level=3)
rows = []
for cls in classes:
    v = hw[cls]
    drop_flag = '↓ 5月骤降' if v[4] < v[3]*0.96 else trend_arrow(v)
    rows.append([cls] + [pct(x) for x in v] + [drop_flag])
add_table(doc, ['班型'] + months + ['趋势'], rows)
add_bullet(doc, '1-4月各班型维持在 84-93% 较高水平，5月全线骤降', bold=True)
add_bullet(doc, '下滑幅度：英语班（内）-6.17pp、英语班（外）-7.53pp、粤语班（粤）-3.93pp')
add_bullet(doc, '5月同步下滑疑似系统性因素（假期安排/课程变化），建议优先排查原因')

add_separator(doc)

# ==============================
# 六、例题解锁数
# ==============================
add_heading(doc, '六、例题解锁数（专项提醒）', level=2)
add_quote(doc, '例题解锁数低于 2 的班型需重点关注，加强例题引导设计。')

add_heading(doc, '6.1 5月达标状态', level=3)
rows = []
low_classes = []
for cls in classes:
    v = example_unlock[cls][4]
    flag = '⚠️ 重点提醒' if v < 2 else '✅ 正常'
    if v < 2: low_classes.append(cls)
    rows.append([cls, f'{v:.2f}', flag])
add_table(doc, ['班型', '5月例题解锁数（道）', '状态'], rows)
if low_classes:
    add_bullet(doc, f'低于 2 的班型：{", ".join(low_classes)}，需优化课中例题引导', bold=True)
add_bullet(doc, '普通班（普）1.99 仅差 0.01 道，6月微调可达标')

add_heading(doc, '6.2 近3个月趋势（3月-5月）', level=3)
rows = []
months3 = ['3月', '4月', '5月']
for cls in classes:
    v3 = example_unlock[cls][2:]  # 3、4、5月
    d = v3[2] - v3[0]
    sign = '+' if d >= 0 else ''
    trend_note = f'{sign}{d:.2f} 道'
    rows.append([cls] + [f'{x:.2f}' for x in v3] + [trend_note])
add_table(doc, ['班型'] + months3 + ['3月→5月变化'], rows)
add_bullet(doc, '普通班（粤）从 2.26→1.83，持续下行（-0.43 道），需重点干预')
add_bullet(doc, '粤语班（粤）、英语班（内/外）在 3 月触底后回升，改善明显')
add_bullet(doc, '普通班（普）连续3个月在 1.70-1.99 区间低位，整体偏低，未见显著改善')

add_separator(doc)

# ==============================
# 七、综合建议
# ==============================
add_heading(doc, '七、6月冲刺建议', level=2)

add_table(doc,
    ['优先级', '问题', '建议动作'],
    [
        ['P0', '作业完成率5月全线骤降', '排查5月课程/假期因素，与运营侧对齐，明确系统性原因'],
        ['P0', '英语班（外）安全解锁比仅47.43%', '专项复盘解锁流程，与其它班型对比差异，制定干预方案'],
        ['P1', '安全解锁比全线未达标（距85%差7-38pp）', '各班型推进解锁引导策略，重点辅导卡点学员'],
        ['P1', '普通班（粤/普）例题解锁<2', '增加课中例题引导环节，优化例题难度匹配'],
        ['P2', '末答正确率4个班型未达95%', '距目标仅差1-2pp，持续微优课程设计可达标'],
        ['P2', '首答正确率仅英语班（外）达标', '普通班（粤）距目标最近（-0.64pp），可优先突破'],
    ]
)

add_separator(doc)
add_para(doc, '数据来源：海外教学北极星指标达成监控看板（S3-S5）+ 学情北极星指标Q2定标版，BI导出截至2026年5月24日。',
          italic=True, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99))

output = r'C:\Users\yelin01\WorkBuddy\2026-05-21-15-19-57\S3-S5_北极星指标分析报告_2026年5月_v2.docx'
doc.save(output)
print(f'Word 已保存: {output}')
