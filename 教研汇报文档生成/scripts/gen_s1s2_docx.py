"""
生成 S1-S2 愉悦度 & 镜头感分析 Word 文档
用法：python gen_s1s2_docx.py <parsed_json路径> <目标图片路径> <输出docx路径>
或者直接用内置数据生成（见 __main__）
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import sys
import os

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table):
    """给表格加边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        borders.append(border)
    tblPr.append(borders)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    return h

def add_para(doc, text, bold=False, italic=False, size=Pt(10.5), color=None, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_bullet(doc, text, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2C3E50')

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r % 2 == 0:
                set_cell_shading(cell, 'F5F7FA')

    doc.add_paragraph()
    return table

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def format_pct(v):
    """将小数转为百分比字符串，保留2位"""
    if v is None:
        return '—'
    try:
        return f"{v*100:.2f}%"
    except:
        return str(v)

def format_diff(v):
    """格式化差距，+Xpp 或 -Xpp"""
    if v is None:
        return '—'
    sign = '+' if v >= 0 else ''
    return f"{sign}{v*100:.2f}pp"

def trend_arrow(vals):
    """根据5个月数据判断趋势"""
    if not vals or len(vals) < 2:
        return '→ 持平'
    first = vals[0]
    last = vals[-1]
    if last > first * 1.05:
        return '↑ 提升'
    elif last < first * 0.95:
        return '↓ 下降'
    else:
        return '→ 持平'

def gen_report(joy_data, lens_data, joy_targets, lens_targets, output_path):
    """
    生成 S1-S2 愉悦度 & 镜头感分析报告
    joy_data: [{"class": "普通班（粤）", "values": [0.036, ...]}, ...]
    lens_data: 同上
    joy_targets: {"普通班（粤）": 0.04, ...}
    lens_targets: {"普通班（粤）": 0.73, ...}
    """
    doc = Document()

    # 样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ========== 标题 ==========
    add_heading(doc, 'S1-S2 学龄前学段愉悦度 & 镜头感分析报告', level=1)
    add_para(doc, '分析时间：2026年5月（数据截至本周）', size=Pt(10))
    add_para(doc, '分析对象：S1-S2 学段（学龄前小班、中班）', size=Pt(10))
    add_para(doc, '核心指标：愉悦度（微笑表情占比）、镜头感（正脸面对屏幕时间占比）', size=Pt(10))
    add_separator(doc)

    # ========== 一、Q2 目标达成 ==========
    add_heading(doc, '一、Q2 目标达成情况（截至5月）', level=2)
    add_quote(doc, 'Q2 指 4月-6月，目标值来自学段季度规划截图。')

    # 愉悦度达标表
    add_heading(doc, '1.1 愉悦度达标分析', level=3)
    joy_rows = []
    for item in joy_data:
        cls = item['class']
        vals = item['values']
        target = joy_targets.get(cls, None)
        actual = vals[-1] if vals else None  # 5月
        if target is not None and actual is not None:
            diff = actual - target
            status = '✅ 已达标' if diff >= 0 else '❌ 未达标'
            joy_rows.append([cls, format_pct(target), format_pct(actual), status, format_diff(diff)])
        elif actual is not None:
            joy_rows.append([cls, '—', format_pct(actual), '—', '—'])
    
    if joy_rows:
        add_table(doc, ['班型', 'Q2 目标', '5月实际', '达成状态', '差距'], joy_rows)
        # 结论（精炼）
        fail_cls = [r[0] for r in joy_rows if '❌' in r[3]]
        if fail_cls:
            add_bullet(doc, f"愉悦度全线未达标，差距最小：{joy_rows[0][0] if joy_rows else '—'}，最大：{fail_cls[-1] if fail_cls else '—'}")
        else:
            add_bullet(doc, "愉悦度全线达标 ✅")

    # 镜头感达标表
    add_heading(doc, '1.2 镜头感达标分析', level=3)
    lens_rows = []
    for item in lens_data:
        cls = item['class']
        vals = item['values']
        target = lens_targets.get(cls, None)
        actual = vals[-1] if vals else None
        if target is not None and actual is not None:
            diff = actual - target
            status = '✅ 已达标' if diff >= 0 else '❌ 未达标'
            lens_rows.append([cls, format_pct(target), format_pct(actual), status, format_diff(diff)])
        elif actual is not None:
            lens_rows.append([cls, '—', format_pct(actual), '—', '—'])

    if lens_rows:
        add_table(doc, ['班型', 'Q2 目标', '5月实际', '达成状态', '差距'], lens_rows)
        pass_cls = [r[0] for r in lens_rows if '✅' in r[3]]
        if pass_cls:
            add_bullet(doc, f"镜头感达标班型：{', '.join(pass_cls)}")
        fail_cls = [r[0] for r in lens_rows if '❌' in r[3]]
        if fail_cls:
            add_bullet(doc, f"未达标：{', '.join(fail_cls)}，差距最大：{lens_rows[[i for i,r in enumerate(lens_rows) if '❌' in r[3]][-1]][0] if fail_cls else '—'}")

    # 达标总览
    add_heading(doc, '1.3 达标情况总览', level=3)
    overview_rows = []
    all_classes = set([item['class'] for item in joy_data] + [item['class'] for item in lens_data])
    for cls in sorted(all_classes):
        joy_item = next((i for i in joy_data if i['class'] == cls), None)
        lens_item = next((i for i in lens_data if i['class'] == cls), None)
        joy_status = '—'
        lens_status = '—'
        if joy_item:
            target = joy_targets.get(cls)
            actual = joy_item['values'][-1] if joy_item['values'] else None
            if target and actual is not None:
                joy_status = '✅' if actual >= target else '❌'
        if lens_item:
            target = lens_targets.get(cls)
            actual = lens_item['values'][-1] if lens_item['values'] else None
            if target and actual is not None:
                lens_status = '✅' if actual >= target else '❌'
        
        status_text = '双项达标' if joy_status == '✅' and lens_status == '✅' else \
                     '单项达标' if joy_status == '✅' or lens_status == '✅' else \
                     '双项未达标' if joy_status == '❌' and lens_status == '❌' else '—'
        overview_rows.append([cls, joy_status, lens_status, status_text])
    
    if overview_rows:
        add_table(doc, ['班型', '愉悦度', '镜头感', '综合状态'], overview_rows)

    add_separator(doc)

    # ========== 二、趋势分析 ==========
    add_heading(doc, '二、1-5月趋势分析', level=2)

    # 愉悦度走势
    add_heading(doc, '2.1 愉悦度走势', level=3)
    month_labels = ['1月', '2月', '3月', '4月', '5月']
    joy_trend_rows = []
    for item in joy_data:
        cls = item['class']
        vals = item['values'][:5]
        trend = trend_arrow(vals)
        joy_trend_rows.append([cls] + [format_pct(v) for v in vals] + [trend])
    
    if joy_trend_rows:
        add_table(doc, ['班型'] + month_labels + ['趋势'], joy_trend_rows)

    # 精炼结论
    # 找进步和退步的班型
    joy_conclusions = []
    for item in joy_data:
        vals = item['values'][:5]
        if len(vals) >= 5:
            if vals[-1] > vals[0] * 1.05:
                joy_conclusions.append(f"{item['class']}↑ 愉悦度回升")
            elif vals[-1] < vals[0] * 0.95:
                joy_conclusions.append(f"{item['class']}↓ 愉悦度下滑")
    if joy_conclusions:
        for c in joy_conclusions:
            add_bullet(doc, c)

    # 镜头感走势
    add_heading(doc, '2.2 镜头感走势', level=3)
    lens_trend_rows = []
    for item in lens_data:
        cls = item['class']
        vals = item['values'][:5]
        trend = trend_arrow(vals)
        lens_trend_rows.append([cls] + [format_pct(v) for v in vals] + [trend])
    
    if lens_trend_rows:
        add_table(doc, ['班型'] + month_labels + ['趋势'], lens_trend_rows)

    # 综合趋势总结
    add_heading(doc, '2.3 综合趋势总结', level=3)
    # 精炼：只写关键发现
    findings = []
    for item in joy_data:
        vals = item['values'][:5]
        if len(vals) >= 5:
            if vals[-1] > vals[0] * 1.03:
                findings.append(f"愉悦度 {item['class']} 近5月整体提升")
            elif vals[-1] < vals[0] * 0.97:
                findings.append(f"愉悦度 {item['class']} 近5月下滑，需关注")
    for item in lens_data:
        vals = item['values'][:5]
        if len(vals) >= 5:
            if vals[-1] > vals[0] * 1.03:
                findings.append(f"镜头感 {item['class']} 近5月持续提升")
            elif vals[-1] < vals[0] * 0.97:
                findings.append(f"镜头感 {item['class']} 近5月走低，需干预")
    
    if findings:
        for f in findings[:5]:  # 最多5条
            add_bullet(doc, f)

    add_separator(doc)

    # ========== 三、问题诊断与建议 ==========
    add_heading(doc, '三、问题诊断与建议', level=2)

    # 优先级表（自动生成）
    priority_rows = []
    # P0: 双项未达标且差距大
    for item in joy_data:
        cls = item['class']
        joy_target = joy_targets.get(cls)
        joy_actual = item['values'][-1] if item['values'] else None
        lens_item = next((i for i in lens_data if i['class'] == cls), None)
        lens_target = lens_targets.get(cls)
        lens_actual = lens_item['values'][-1] if lens_item and lens_item['values'] else None
        
        joy_gap = (joy_actual - joy_target) if (joy_target and joy_actual) else 0
        lens_gap = (lens_actual - lens_target) if (lens_target and lens_actual) else 0
        
        both_fail = (joy_target and joy_actual and joy_actual < joy_target) and \
                   (lens_target and lens_actual and lens_actual < lens_target)
        one_fail = (joy_target and joy_actual and joy_actual < joy_target) or \
                  (lens_target and lens_actual and lens_actual < lens_target)
        
        if both_fail:
            priority_rows.append(['P0', cls, '双项未达标', '教研侧介入，排查课堂设计'])
        elif one_fail:
            priority_rows.append(['P1', cls, '单项未达标', '微调互动设计，争取下月达标'])
        else:
            priority_rows.append(['P2', cls, '已达标', '总结经验，维持稳定'])
    
    if priority_rows:
        add_table(doc, ['优先级', '班型', '核心问题', '建议动作'], priority_rows)

    # 6月冲刺建议（精炼，最多3条）
    add_heading(doc, '3.1 6月冲刺建议', level=3)
    
    # 找最接近达标的
    closest_joy = None
    closest_joy_gap = float('inf')
    for item in joy_data:
        cls = item['class']
        target = joy_targets.get(cls)
        actual = item['values'][-1] if item['values'] else None
        if target and actual and actual < target:
            gap = target - actual
            if gap < closest_joy_gap:
                closest_joy_gap = gap
                closest_joy = cls
    
    closest_lens = None
    closest_lens_gap = float('inf')
    for item in lens_data:
        cls = item['class']
        target = lens_targets.get(cls)
        actual = item['values'][-1] if item['values'] else None
        if target and actual and actual < target:
            gap = target - actual
            if gap < closest_lens_gap:
                closest_lens_gap = gap
                closest_lens = cls
    
    if closest_joy:
        add_bullet(doc, f"{closest_joy} 愉悦度距目标仅 {closest_joy_gap*100:.2f}pp，重点突破")
    if closest_lens:
        add_bullet(doc, f"{closest_lens} 镜头感距目标仅 {closest_lens_gap*100:.2f}pp，争取达标")
    
    # 已达标的
    passed = [r[0] for r in (priority_rows or []) if 'P2' in r[0]]
    if passed:
        add_bullet(doc, f"{passed[0]} 已达标，维持稳定并提炼经验")

    add_separator(doc)

    # 页脚
    add_para(doc, '数据来源：海外教学北极星指标达成监控看板（S1-S2），BI导出数据截至2026年5月。', 
              italic=True, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99))

    doc.save(output_path)
    print(f"Word 文档已保存: {output_path}")

if __name__ == '__main__':
    # 内置示例数据（当用户没有提供解析JSON时使用）
    # 替换为实际解析到的数据
    joy_data = [
        {"class": "普通班（粤）", "values": [0.0363, 0.0394, 0.0392, 0.0374, 0.0330]},
        {"class": "普通班（普）", "values": [0.0186, 0.0198, 0.0207, 0.0177, 0.0176]},
        {"class": "粤语班（粤）", "values": [0.0233, 0.0237, 0.0214, 0.0236, 0.0234]},
        {"class": "英语班（内）", "values": [0.0233, 0.0231, 0.0213, 0.0215, 0.0215]},
        {"class": "英语班（外）", "values": [0.0282, 0.0268, 0.0234, 0.0246, 0.0248]},
    ]
    lens_data = [
        {"class": "普通班（粤）", "values": [0.7097, 0.7053, 0.7188, 0.7181, 0.7083]},
        {"class": "普通班（普）", "values": [0.6398, 0.6310, 0.6351, 0.6392, 0.6161]},
        {"class": "粤语班（粤）", "values": [0.6844, 0.6849, 0.6804, 0.6902, 0.6942]},
        {"class": "英语班（内）", "values": [0.6117, 0.6226, 0.6194, 0.6093, 0.6081]},
        {"class": "英语班（外）", "values": [0.6323, 0.6283, 0.6497, 0.6620, 0.6518]},
    ]
    joy_targets = {
        "普通班（粤）": 0.040,
        "普通班（普）": 0.025,
        "粤语班（粤）": 0.025,
        "英语班（内）": 0.025,
        "英语班（外）": 0.030,
    }
    lens_targets = {
        "普通班（粤）": 0.73,
        "普通班（普）": 0.65,
        "粤语班（粤）": 0.70,
        "英语班（内）": 0.65,
        "英语班（外）": 0.65,
    }

    output = r"C:\Users\yelin01\WorkBuddy\2026-05-21-15-19-57\S1-S2_愉悦度镜头感分析报告_2026年5月.docx"
    gen_report(joy_data, lens_data, joy_targets, lens_targets, output)
