"""
生成 Word 汇报文档 - 教研汇报技能
用法：python gen_report_docx.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table):
    """给表格加边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        borders.append(border)
    tblPr.append(borders)


def add_heading(doc, text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(10.5), color=None, alignment=None, space_after=Pt(6)):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    for run in p.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10.5)
    return p


def add_quote(doc, text):
    """添加引用块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True
    return p


def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2C3E50')
    
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r % 2 == 0:
                set_cell_shading(cell, 'F5F7FA')
    
    doc.add_paragraph()  # 表后空行
    return table


def add_separator(doc):
    """添加分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_report(data_dict, output_path):
    """
    生成汇报文档
    data_dict: 从 Excel/Markdown 解析后的数据字典
    包含: indicators, v9_progress, org_hr, business_support, subject_work, plans 等
    """
    doc = Document()
    
    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    
    # 样式设置
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ===== 标题 =====
    add_heading(doc, '海外教研工作汇报', level=1)
    add_para(doc, '汇报时间：2026年5月（截至0519）', size=Pt(10))
    add_para(doc, '汇报人：海外教研团队', size=Pt(10))
    add_separator(doc)
    
    # ===== 一、北极星指标 =====
    add_heading(doc, '一、北极星指标达成情况', level=2)
    add_quote(doc, '统计口径：2025年11月 – 2026年5月，覆盖 S1-S2、S3-S5、S6-S7 三个学段')
    
    add_heading(doc, '1.1 核心"四率"总览', level=3)
    # 这里用 data_dict['indicators'] 填充表格
    # 示例：
    add_table(doc,
        ['指标', '学段', '普通班（粤）', '普通班（普）', '粤语班（粤）', '英语班（内）', '英语班（外）', '目标'],
        data_dict.get('indicators', [])
    )
    
    # ===== 二、V9研发 =====
    add_heading(doc, '二、V9课程研发进度', level=2)
    # 填充 data_dict['v9_progress']...
    
    # ===== 三、组织人力 =====
    add_heading(doc, '三、组织与人力', level=2)
    # 填充 data_dict['org_hr']...
    
    # ===== 四、业务支持 =====
    add_heading(doc, '四、4-5月业务支持输出', level=2)
    # 填充 data_dict['business_support']...
    
    # ===== 五、学科工作 =====
    add_heading(doc, '五、学科工作进展', level=2)
    # 填充 data_dict['subject_work']...
    
    # ===== 六、重点专项 =====
    add_heading(doc, '六、重点专项', level=2)
    # 填充 data_dict['key_projects']...
    
    # ===== 七、月底计划 =====
    add_heading(doc, '七、月底工作计划', level=2)
    add_table(doc,
        ['计划事项', '优先级', '负责方向'],
        data_dict.get('plans', [])
    )
    
    # 保存
    doc.save(output_path)
    print(f"Word文档已保存: {output_path}")


if __name__ == "__main__":
    # 测试用 - 实际由调用方传入 data_dict
    test_data = {
        'indicators': [
            ['课中首答正确率', 'S6-S7', '79.1%', '75.7%', '76.6%', '76.1%', '73.6%', '77%'],
            ['课中末答正确率', 'S6-S7', '95.9%', '93.7%', '95.1%', '92.6%', '96.3%', '95%'],
        ],
        'plans': [
            ['教研绩效组内宣讲及落地执行', 'P0', '人力'],
            ['课程研发持续推进', 'P0', '研发'],
        ]
    }
    
    output = r"C:\Users\yelin01\WorkBuddy\test_report.docx"
    generate_report(test_data, output)
